# -*- coding: utf-8 -*-
"""
Build the dissertation as a .docx conforming to the 1400 writing guideline:
- A4, margins R3 / L3 / T3.5 / B3 cm, line spacing 1.3, RTL + justified
- Fonts: B Nazanin (Persian, w:cs) + Times New Roman (English, w:ascii/hAnsi)
- Body 14, chapter titles bold 14, sub-titles bold 12, FA abstract 12,
  EN abstract Times 11, references 11, footnotes 11
- Real Word footnotes from {{fn:N}} markers, clickable TOC field
- Cover, بسم‌الله, dedication, FA abstract, TOC, chapters (each new page),
  glossary, debate, methodology, references (FA then Latin), EN abstract
Source of truth: web/thesis.json
"""
import os, re, json
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.part import Part
from docx.opc.packuri import PackURI

HERE = os.path.dirname(os.path.abspath(__file__))
# Prefer the enriched thesis (full reference list, fixed footnotes, expanded
# chapters 4 & 5). Fall back to the base thesis.json if enrichment is absent.
_ENRICHED = os.path.join(HERE, "..", "web", "thesis_enriched.json")
TH_PATH = _ENRICHED if os.path.exists(_ENRICHED) else os.path.join(HERE, "..", "web", "thesis.json")
OUT = os.path.join(HERE, "..", "web", "thesis_draft.docx")

FA_FONT = "B Nazanin"
EN_FONT = "Times New Roman"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def w(tag): return qn("w:" + tag)

def set_run_fonts(run, fa_pt, en_pt, bold=False, color=None):
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(w("rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.insert(0, rfonts)
    rfonts.set(w("ascii"), EN_FONT); rfonts.set(w("hAnsi"), EN_FONT); rfonts.set(w("cs"), FA_FONT)
    sz = OxmlElement("w:sz"); sz.set(w("val"), str(int(en_pt*2))); rpr.append(sz)
    szcs = OxmlElement("w:szCs"); szcs.set(w("val"), str(int(fa_pt*2))); rpr.append(szcs)
    rtl = OxmlElement("w:rtl"); rtl.set(w("val"), "1"); rpr.append(rtl)
    if bold:
        rpr.append(OxmlElement("w:b")); rpr.append(OxmlElement("w:bCs"))
    if color is not None:
        c = OxmlElement("w:color"); c.set(w("val"), color); rpr.append(c)

def set_run_fonts_ltr(run, en_pt, bold=False):
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(w("rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.insert(0, rfonts)
    rfonts.set(w("ascii"), EN_FONT); rfonts.set(w("hAnsi"), EN_FONT); rfonts.set(w("cs"), EN_FONT)
    sz = OxmlElement("w:sz"); sz.set(w("val"), str(int(en_pt*2))); rpr.append(sz)
    szcs = OxmlElement("w:szCs"); szcs.set(w("val"), str(int(en_pt*2))); rpr.append(szcs)
    if bold:
        rpr.append(OxmlElement("w:b")); rpr.append(OxmlElement("w:bCs"))

def make_rtl(paragraph, justify=True, line=1.3, before=0, after=6):
    ppr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi"); bidi.set(w("val"), "1"); ppr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)

# ---- footnotes ----
class Footnotes:
    def __init__(self, doc):
        self.doc = doc; self.items = []; self._next_id = 1
    def add(self, paragraph, text):
        fid = self._next_id; self._next_id += 1
        self.items.append((fid, text))
        run = paragraph.add_run()
        rpr = run._r.get_or_add_rPr()
        rstyle = OxmlElement("w:rStyle"); rstyle.set(w("val"), "FootnoteReference"); rpr.append(rstyle)
        rtl = OxmlElement("w:rtl"); rtl.set(w("val"), "1"); rpr.append(rtl)
        ref = OxmlElement("w:footnoteReference"); ref.set(w("id"), str(fid))
        run._r.append(ref)
        return fid
    def _xml(self, fid, text):
        esc = text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        return (f'<w:footnote w:id="{fid}"><w:p><w:pPr><w:pStyle w:val="FootnoteText"/>'
                f'<w:bidi w:val="1"/><w:jc w:val="both"/></w:pPr>'
                f'<w:r><w:rPr><w:rStyle w:val="FootnoteReference"/><w:rtl w:val="1"/></w:rPr>'
                f'<w:footnoteRef/></w:r>'
                f'<w:r><w:rPr><w:rFonts w:ascii="{EN_FONT}" w:hAnsi="{EN_FONT}" w:cs="{FA_FONT}"/>'
                f'<w:sz w:val="22"/><w:szCs w:val="22"/><w:rtl w:val="1"/></w:rPr>'
                f'<w:t xml:space="preserve"> {esc}</w:t></w:r></w:p></w:footnote>')
    def finalize(self):
        parts = ['<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
                 f'<w:footnotes xmlns:w="{W}">',
                 '<w:footnote w:type="separator" w:id="-1"><w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p></w:footnote>',
                 '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>']
        for fid, text in self.items:
            parts.append(self._xml(fid, text))
        parts.append('</w:footnotes>')
        xml = "".join(parts).encode("utf-8")
        partname = PackURI("/word/footnotes.xml")
        ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
        fn_part = Part(partname, ct, xml, self.doc.part.package)
        reltype = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
        self.doc.part.relate_to(fn_part, reltype)

def build_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = EN_FONT; normal.font.size = Pt(14)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(w("rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.insert(0, rfonts)
    rfonts.set(w("ascii"), EN_FONT); rfonts.set(w("hAnsi"), EN_FONT); rfonts.set(w("cs"), FA_FONT)
    szcs = OxmlElement("w:szCs"); szcs.set(w("val"), "28"); rpr.append(szcs)
    names = [s.name for s in styles]
    if "FootnoteText" not in names:
        fts = styles.add_style("FootnoteText", WD_STYLE_TYPE.PARAGRAPH)
        fts.font.name = EN_FONT; fts.font.size = Pt(11)
        fts.paragraph_format.space_after = Pt(0)
    if "FootnoteReference" not in names:
        frs = styles.add_style("FootnoteReference", WD_STYLE_TYPE.CHARACTER)
        frs.font.size = Pt(11)
        va = OxmlElement("w:vertAlign"); va.set(w("val"), "superscript")
        frs.element.get_or_add_rPr().append(va)
    # Custom RTL heading styles. Critically they carry NO outlineLvl: any
    # paragraph (or style) with an outline level is remapped by LibreOffice
    # onto its own LTR built-in Heading style, which drops our bidi and forces
    # the heading left-to-right. Without outlineLvl the paragraph stays RTL,
    # and the TOC field collects these headings by STYLE NAME via the \t switch.
    for nm in ("GPEHeading1", "GPEHeading2", "GPEHeading3"):
        if nm in names:
            continue
        st = styles.add_style(nm, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = styles["Normal"]
        st.next_paragraph_style = styles["Normal"]
        ppr = st.element.get_or_add_pPr()
        kn = OxmlElement("w:keepNext"); ppr.append(kn)
        bd = OxmlElement("w:bidi"); bd.set(w("val"), "1"); ppr.append(bd)
        jc = OxmlElement("w:jc"); jc.set(w("val"), "right"); ppr.append(jc)
        rpr = st.element.get_or_add_rPr()
        rt = OxmlElement("w:rtl"); rt.set(w("val"), "1"); rpr.append(rt)
        rpr.append(OxmlElement("w:b")); rpr.append(OxmlElement("w:bCs"))


FN_RE = re.compile(r"\{\{fn:(\d+)\}\}")

def add_para(doc, fns, text, fa_pt=14, en_pt=11, bold=False, justify=True,
             line=1.3, before=0, after=6, indent=None, section_fns=None):
    p = doc.add_paragraph(); make_rtl(p, justify=justify, line=line, before=before, after=after)
    if indent:
        p.paragraph_format.first_line_indent = Mm(indent)
    pos = 0
    for m in FN_RE.finditer(text):
        pre = text[pos:m.start()]
        if pre:
            r = p.add_run(pre); set_run_fonts(r, fa_pt, en_pt, bold=bold)
        n = int(m.group(1))
        fn_text = ""
        if section_fns and 1 <= n <= len(section_fns):
            fn_text = section_fns[n-1]
        fns.add(p, fn_text)
        pos = m.end()
    tail = text[pos:]
    if tail:
        r = p.add_run(tail); set_run_fonts(r, fa_pt, en_pt, bold=bold)
    return p

def add_heading(doc, num, title, level, fa_pt):
    # Plain RTL paragraph (NOT built-in Heading styles). LibreOffice merges the
    # built-in "Heading N" style names with its own LTR built-ins and drops our
    # bidi/jc, rendering headings left-to-right. A plain bidi paragraph with a
    # direct outlineLvl renders RTL correctly AND is still collected by the TOC
    # field (\o "1-3" gathers paragraphs by outline level, not by style name).
    nm = {1: "GPEHeading1", 2: "GPEHeading2", 3: "GPEHeading3"}.get(level, "GPEHeading3")
    p = doc.add_paragraph(style=doc.styles[nm])
    make_rtl(p, justify=False, before=(14 if level == 1 else 10), after=6)
    txt = (f"{num} " if num else "") + title
    r = p.add_run(txt)
    set_run_fonts(r, fa_pt, 12, bold=True, color="14203a")
    return p

def center_para(doc, text, fa_pt=14, bold=False, before=0, after=6):
    p = doc.add_paragraph(); make_rtl(p, justify=False, before=before, after=after)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); set_run_fonts(r, fa_pt, 12, bold=bold)
    return p

def add_toc(doc):
    p = doc.add_paragraph(); make_rtl(p, justify=False); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(w("fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\h \\z \\t "GPEHeading1,1,GPEHeading2,2,GPEHeading3,3"'
    sep = OxmlElement("w:fldChar"); sep.set(w("fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "برای به‌روزرسانی فهرست، کلید F9 را فشار دهید."
    e = OxmlElement("w:fldChar"); e.set(w("fldCharType"), "end")
    for x in (b, instr, sep, t, e): run._r.append(x)

def enable_update_fields(doc):
    uf = OxmlElement("w:updateFields"); uf.set(w("val"), "true")
    doc.settings.element.append(uf)

def set_rtl_section(doc):
    for section in doc.sections:
        section._sectPr.append(OxmlElement("w:bidi"))

def main():
    TH = json.load(open(TH_PATH, encoding="utf-8"))
    m = TH.get("meta", {})
    ab = TH.get("en_abstract", {})
    doc = Document(); build_styles(doc); fns = Footnotes(doc)

    sec = doc.sections[0]
    sec.page_height = Mm(297); sec.page_width = Mm(210)
    sec.top_margin = Mm(35); sec.bottom_margin = Mm(30)
    sec.right_margin = Mm(30); sec.left_margin = Mm(30)

    # ── Cover ──
    center_para(doc, m.get("university_fa",""), fa_pt=16, bold=True, before=24, after=6)
    center_para(doc, m.get("degree_fa",""), fa_pt=13, after=30)
    center_para(doc, "عنوان:", fa_pt=13, bold=True, after=6)
    center_para(doc, m.get("subtitle_fa", m.get("title_fa","")), fa_pt=18, bold=True, after=10)
    p = doc.add_paragraph(); make_rtl(p, justify=False, after=30); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(ab.get("title_en","")); set_run_fonts_ltr(r, 12, bold=False)
    center_para(doc, m.get("supervisor_fa",""), fa_pt=14, bold=True, after=6)
    center_para(doc, "سال تحصیلی ۱۴۰۴–۱۴۰۵", fa_pt=13, after=6)
    doc.add_page_break()

    # ── بسم‌الله ──
    center_para(doc, "بِسْمِ اللَّهِ الرَّحْمَٰنِ الرَّحِيمِ", fa_pt=16, bold=True, before=40, after=40)
    doc.add_page_break()

    # ── Dedication ──
    center_para(doc, "تقدیم", fa_pt=14, bold=True, before=30, after=10)
    center_para(doc, "این بخش (تقدیم‌نامه) مطابق سلیقهٔ نگارنده تکمیل می‌شود.", fa_pt=13, after=6)
    doc.add_page_break()

    # ── FA abstract ──
    p = doc.add_paragraph(); make_rtl(p, justify=False, after=8); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("چکیده"); set_run_fonts(r, 12, 11, bold=True)
    fa_abs = ("این رساله فرایندِ شکل‌گیریِ «پیمان جهانی محیط زیست» (۲۰۱۷–۲۰۲۲) و تأثیرِ آن بر موافقت‌نامه‌های چندجانبهٔ محیط‌زیستی (MEAs) را بررسی می‌کند. "
              "با روشِ توصیفی‑تحلیلیِ حقوقی و با تکیه بر سندِ چندپارگیِ کمیسیون حقوق بین‌الملل (۲۰۰۶) به‌عنوان چارچوبِ نظری، استدلال می‌شود که هدفِ پیمان — گردآوریِ اصولِ بنیادین در یک سندِ چترِ الزام‌آور و گذارِ تدریجیِ اصول از حقوق نرم به حقوق سخت از رهگذرِ «یکپارچه‌سازیِ نظام‌مند» (مادهٔ ۳۱(۳)(ج) کنوانسیون وین) — موجه بود، اما ابزارِ برگزیده (معاهدهٔ فراگیرِ اجماع‌محورِ واحد) با سرشتِ واکنشی، بخشی و تدریجیِ حقوق بین‌الملل محیط زیست ناسازگار افتاد و مذاکرات در ۲۰۲۲ به معاهده نینجامید. "
              "یافتهٔ کانونی این است که تأثیرِ واقعیِ پیمان نه از تصویبِ آن، بلکه از فرایندِ شکل‌گیری و ناکامیِ آن برخاست: تقویتِ گفتمانِ انسجام‌بخشی و زمینه‌سازی برای شناساییِ حقِ جهانی بر محیط‌زیستِ سالم در قطعنامهٔ ۷۶/۳۰۰ مجمع عمومی. "
              "رساله در پایان نقشهٔ راهی شش‌ستونه برای تحققِ همان هدف از مسیرهای واقع‌بینانه‌تر (توسعهٔ قضایی، تقویتِ نهادیِ برنامهٔ محیط زیست ملل متحد، بازنگریِ دوره‌ای، خوشه‌بندیِ معاهدات، کاربستِ راهبردیِ حقوق نرم و تفسیرِ یکسانِ اصول) پیشنهاد می‌کند و دلالت‌های آن را برای ایران بازمی‌نماید.")
    add_para(doc, fns, fa_abs, fa_pt=12, en_pt=11, after=8)
    p = doc.add_paragraph(); make_rtl(p, after=6)
    r = p.add_run("واژگان کلیدی: "); set_run_fonts(r, 12, 11, bold=True)
    r2 = p.add_run("پیمان جهانی محیط زیست، موافقت‌نامه‌های چندجانبهٔ محیط‌زیستی، چندپارگیِ حقوق بین‌الملل، یکپارچه‌سازیِ نظام‌مند، حق بر محیط‌زیستِ سالم، سند چتر.")
    set_run_fonts(r2, 12, 11)
    doc.add_page_break()

    # ── TOC ──
    p = doc.add_paragraph(); make_rtl(p, justify=False, after=6); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("فهرست مطالب"); set_run_fonts(r, 14, 12, bold=True)
    add_toc(doc)
    doc.add_page_break()

    # ── Chapters ──
    for c in TH.get("chapters", []):
        add_heading(doc, "", f"{c.get('num','')}: {c.get('title_fa','')}", 1, 14)
        if c.get("summary_fa"):
            add_para(doc, fns, c["summary_fa"], fa_pt=12, en_pt=11, after=8)
        for s in c.get("sections", []):
            lvl = min(s.get("level",1), 3)
            hlvl = 2 if lvl == 1 else 3
            add_heading(doc, s.get("num",""), s.get("title_fa",""), hlvl, 12)
            sfns = s.get("fns", [])
            for para in s.get("paras", []):
                add_para(doc, fns, para, fa_pt=14, en_pt=11, after=6, indent=8, section_fns=sfns)
        doc.add_page_break()

    # ── Glossary ──
    gl = TH.get("glossary", [])
    if gl:
        add_heading(doc, "", "واژه‌نامهٔ مستندِ اصطلاحات", 1, 14)
        for g in gl:
            p = doc.add_paragraph(); make_rtl(p, after=4)
            r = p.add_run(f"{g.get('term_fa','')} ({g.get('term_en','')}): ")
            set_run_fonts(r, 13, 11, bold=True)
            r2 = p.add_run(g.get("def_fa","")); set_run_fonts(r2, 13, 11)
            add_para(doc, fns, "منبع: " + g.get("source_fa",""), fa_pt=11, en_pt=10, after=8, indent=6)
        doc.add_page_break()

    # ── Debate ──
    db = TH.get("debate", {})
    if db.get("proponents") or db.get("opponents"):
        add_heading(doc, "", "موافقان و مخالفانِ پیمان جهانی", 1, 14)
        if db.get("proponents"):
            add_heading(doc, "", "دیدگاهِ موافقان", 2, 12)
            for x in db["proponents"]:
                add_para(doc, fns, f"• {x.get('point_fa','')} (منبع: {x.get('source_fa','')})", fa_pt=13, en_pt=11, after=4)
        if db.get("opponents"):
            add_heading(doc, "", "دیدگاهِ مخالفان", 2, 12)
            for x in db["opponents"]:
                add_para(doc, fns, f"• {x.get('point_fa','')} (منبع: {x.get('source_fa','')})", fa_pt=13, en_pt=11, after=4)
        if db.get("synthesis_fa"):
            add_heading(doc, "", "جمع‌بندیِ داوری‌شده", 2, 12)
            add_para(doc, fns, db["synthesis_fa"], fa_pt=13, en_pt=11, after=8)
        doc.add_page_break()

    # ── Methodology ──
    cm = TH.get("conclusion_methodology", {})
    if cm.get("models"):
        add_heading(doc, "", "روش‌شناسیِ نتیجه‌گیری", 1, 14)
        if cm.get("intro_fa"):
            add_para(doc, fns, cm["intro_fa"], fa_pt=13, en_pt=11, after=8)
        for md in cm["models"]:
            add_heading(doc, "", md.get("name_fa",""), 2, 12)
            add_para(doc, fns, md.get("approach_fa",""), fa_pt=13, en_pt=11, after=4)
            add_para(doc, fns, "کاربست در این رساله: " + md.get("adopt_fa",""), fa_pt=13, en_pt=11, after=4)
            add_para(doc, fns, "منبع: " + md.get("source_fa",""), fa_pt=11, en_pt=10, after=8, indent=6)
        if cm.get("our_method_fa"):
            add_heading(doc, "", "روشِ برگزیدهٔ این رساله", 2, 12)
            add_para(doc, fns, cm["our_method_fa"], fa_pt=13, en_pt=11, after=8)
        doc.add_page_break()

    # ── References ──
    refs = TH.get("references", {})
    add_heading(doc, "", "فهرست منابع و مآخذ", 1, 14)
    add_heading(doc, "", "الف) منابع فارسی", 2, 12)
    for r0 in refs.get("fa", []):
        add_para(doc, fns, r0.get("text",""), fa_pt=11, en_pt=11, after=4)
    add_heading(doc, "", "ب) منابع لاتین (Latin References)", 2, 12)
    for r0 in refs.get("en", []):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.3; p.paragraph_format.space_after = Pt(4)
        rr = p.add_run(r0.get("text","")); set_run_fonts_ltr(rr, 11)
    doc.add_page_break()

    # ── EN abstract ──
    if ab.get("body_en"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
        r = p.add_run("Abstract"); set_run_fonts_ltr(r, 12, bold=True)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
        r = p.add_run(ab.get("title_en","")); set_run_fonts_ltr(r, 12, bold=True)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.3
        r = p.add_run(ab["body_en"]); set_run_fonts_ltr(r, 11)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_before = Pt(8)
        r = p.add_run("Keywords: "); set_run_fonts_ltr(r, 11, bold=True)
        r2 = p.add_run("; ".join(ab.get("keywords_en", [])) + "."); set_run_fonts_ltr(r2, 11)

    # heading styles → RTL + B Nazanin
    for hname in ("Heading 1","Heading 2","Heading 3"):
        st = doc.styles[hname]
        st.font.name = EN_FONT; st.font.color.rgb = RGBColor(0,0,0)
        rpr = st.element.get_or_add_rPr()
        rf = rpr.find(w("rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
        rf.set(w("cs"), FA_FONT); rf.set(w("ascii"), EN_FONT); rf.set(w("hAnsi"), EN_FONT)

    fns.finalize(); set_rtl_section(doc); enable_update_fields(doc)
    doc.save(OUT)
    print("saved:", OUT, os.path.getsize(OUT), "bytes; footnotes:", len(fns.items))

if __name__ == "__main__":
    main()
