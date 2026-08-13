# -*- coding: utf-8 -*-
"""
Build a comprehensive, professional Persian report documenting the full
digest of the key uploaded source PDFs: what was extracted, where/how each
was used in the dissertation, and how each optimized the conclusion/scenario,
plus the concrete conclusion upgrades applied after this digest.

Reuses the RTL / B-Nazanin machinery from build_thesis_docx.py.
Source of truth: web/db_mining.json (DB35-DB40) + web/thesis.json.
Output: web/گزارش_تحلیل_منابع.docx  (+ PDF via libreoffice)
"""
import os, json
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import build_thesis_docx as B  # reuse helpers

# ── brand palette (matches the site: deep petrol + gold on light) ──
NAVY   = "123a52"   # H1 bars
GOLD   = "b08d3f"   # accents / rules
INK    = "14203a"   # strong text
SUBBG  = "eaf1f6"   # light subheading band
CARDBG = "f5f8fb"   # bibliographic card
GREEN  = "1e6b3a"
RED    = "b03a2e"
PURPLE = "5a4a8a"

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "..", "web", "گزارش_تحلیل_منابع.docx")
w = B.w

INTRO = [
    "این گزارش، حاصلِ «هضمِ» تحلیلیِ کاملِ منابعِ کلیدیِ آپلودی در سامانهٔ هوشمندِ پژوهش است. "
    "هر منبع به‌صورتِ مستقل واکاوی، نکاتِ راهبردیِ آن استخراج، و جایگاهِ دقیقِ آن در بدنهٔ رساله "
    "و در فرایندِ نتیجه‌گیری مشخص شده است. هدفِ این سند، شفاف‌سازیِ زنجیرهٔ استناد است: "
    "اینکه هر یافته از کجا آمده، در کدام فصل و کدام پرسشِ پژوهش به‌کار رفته، و چگونه به بهینه‌سازیِ "
    "نتیجه‌گیری در «سناریوی رسالهٔ ما» یاری رسانده است.",

    "روشِ کار بدین‌گونه بوده است: نخست متنِ کاملِ هر منبع پردازش و دایجست شده؛ سپس یافتهٔ اصلی، "
    "نقاطِ قوت و ضعفِ روش‌شناختیِ منبع، و «یادداشتِ استناد» (اینکه در کدام فصل و ذیلِ کدام پرسش "
    "به‌کار می‌آید) ثبت شده است. منابعِ داوری‌نشده (پیش‌چاپ) به‌صراحت علامت‌گذاری شده‌اند تا "
    "اعتبارِ علمیِ استناد خدشه‌دار نشود. آنچه در پی می‌آید، برای هر منبع شاملِ چهار بخش است: "
    "(۱) شناسنامهٔ کتاب‌شناختی، (۲) دایجستِ کامل و نکاتِ استخراج‌شده، (۳) کجا و چگونه در پژوهش "
    "به‌کار رفت، و (۴) سهمِ آن در بهینه‌سازیِ نتیجه‌گیری و سناریو. در پایان نیز بخشی مستقل به "
    "«ارتقاهای عملیِ اعمال‌شده بر نتیجه‌گیری در پیِ این هضم» اختصاص یافته است.",
]

ORDER = ["DB35", "DB39", "DB38", "DB37", "DB36", "DB40"]

# Persian-first display identity for each source so that headings and the
# reference list read RTL/Persian-first (never lead with an English string).
# The English author string / title is still shown, but only as a labelled
# inline value inside the bibliographic identity block.
PERSIAN = {
    "DB35": {"name": "تیگره",
             "title": "تحولِ حقوق بین‌الملل محیط زیست در میانهٔ بن‌بستِ سیاسی: حقوقِ محیط‌زیستی همچون زمینهٔ مشترک"},
    "DB39": {"name": "تیلر",
             "title": "آیا اهدافِ پیمانِ جهانیِ پیشنهادیِ محیط زیست مطلوب‌اند و آیا پیمان ارزش‌افزوده‌ای به حقوق بین‌الملل محیط زیست می‌افزاید؟"},
    "DB38": {"name": "چاباسون و اِژه",
             "title": "شکستِ پیمانِ جهانیِ محیط زیست: فرصتی ازدست‌رفته یا تیری ازکنارگذشته؟"},
    "DB37": {"name": "کمیسیون حقوق بین‌الملل (کوسکِنیِمی)",
             "title": "نتیجه‌گیریِ کارِ گروهِ مطالعاتیِ چندپارگیِ حقوق بین‌الملل: دشواری‌های ناشی از تنوع و گسترشِ حقوق بین‌الملل"},
    "DB36": {"name": "داموآ و همکاران",
             "title": "بن‌بست در انتروپوسن: معاهداتِ ناکارآمد و کشمکش بر سرِ حکمرانیِ سیّاره‌ای"},
    "DB40": {"name": "پیش‌نویسِ پیمان (گروهِ حقوق‌دانان / IUCN)",
             "title": "پیش‌نویسِ پیمانِ جهانیِ محیط زیست — ۲۶ ماده"},
}

_FA_DIGITS = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")
def fa_num(x):
    return str(x).translate(_FA_DIGITS)

NARR = {
    "DB35": {
        "points": [
            "در شرایطِ بن‌بستِ سیاسیِ حقوق بین‌الملل محیط زیست، «حقوقِ محیط‌زیستی» — به‌ویژه حق بر محیط‌زیستِ سالم — می‌تواند زمینهٔ مشترکِ پیشبردِ تدریجیِ نظام باشد.",
            "توسعهٔ حقوق نه از راهِ یک معاهدهٔ فراگیرِ واحد، بلکه از مسیرِ رویه، آرای قضایی و شناساییِ تدریجیِ حقوق پیش می‌رود.",
            "الگوی روش‌شناختیِ نتیجه‌گیری سه‌لایه است: سنتزِ یافته‌ها ← طرحِ گزینه‌های پیشِ‌رو ← چشم‌اندازِ دگرگونیِ بوم‌شناختی، با استنادِ متراکمِ درون‌متنی.",
            "پیوندِ مستقیمِ فرایندِ پیمان با قطعنامهٔ ۷۶/۳۰۰ مجمع عمومی (حق بر محیط‌زیستِ سالم) به‌مثابهٔ «دستاوردِ پنهانِ» ابتکار.",
        ],
        "usage": "این رساله، به‌سببِ روش‌شناسیِ دقیق و پرارجاعِ خود، منبعِ طلاییِ الگوی نتیجه‌گیری بوده است. "
                 "یافته‌های آن در فصلِ دوم (سیر تحولِ حقوق و ظهورِ گفتمانِ حق‌محور)، فصلِ سوم (تحلیلِ اصولِ پیمان)، "
                 "فصلِ چهارم (پیامدهای شکست و مسیرِ جایگزین) و فصلِ پنجم (الگوی نتیجه‌گیری) به‌کار رفته و به پرسشِ اصلی "
                 "و دو پرسشِ فرعیِ نخست پاسخ می‌دهد.",
        "opt": "الگوی «استنادِ متراکمِ درون‌متنی» و ساختارِ «سنتز ← گزینه‌ها» مستقیماً از تیگره وام گرفته شد و "
               "به نتیجه‌گیریِ رساله سطحِ استانداردِ دکتری بخشید. مهم‌تر آنکه این منبع کمک کرد نتیجه‌گیریِ ما از "
               "«روایتِ شکستِ صِرف» فراتر رود و به تحلیلِ «دستاوردهای پنهانِ فرایندِ پیمان» (تثبیتِ حق بر محیط‌زیستِ سالم "
               "در قطعنامهٔ ۷۶/۳۰۰) برسد. ستونِ یکم و پنجمِ نقشهٔ راهِ جایگزین (توسعهٔ رویهٔ قضایی و حقوق نرم) مستقیماً بر تزِ این منبع استوار شده‌اند.",
    },
    "DB39": {
        "points": [
            "اهدافِ پیمان (تحکیم و انسجام‌بخشیِ اصول) مطلوب‌اند؛ اما ارزش‌افزودهٔ عملیِ آن محدود است.",
            "علتِ محدودیت: ضعفِ سازوکارِ اجرا و خطرِ تضعیفِ رژیم‌های تخصصیِ موجود (MEAها).",
            "داوریِ نهایی به‌صورتِ گزارهٔ متوازن: «هدف درست است، اما ابزار محلِ تردید».",
            "پرسشِ پژوهشِ این اثر (مطلوبیتِ اهداف + سنجشِ ارزش‌افزوده) تقریباً منطبق بر پرسش‌های رسالهٔ ماست.",
        ],
        "usage": "این پژوهش در فصلِ سوم (گفتارِ «مطلوبیتِ اهداف و ارزش‌افزوده»)، فصلِ چهارم (اثرِ پیمان بر MEAها) و "
                 "فصلِ پنجم (داوریِ متوازنِ نتیجه‌گیری) به‌کار رفت و به هر سه پرسشِ پژوهش مربوط می‌شود. از آنجا که "
                 "پرسشِ محوریِ آن آینهٔ پرسشِ رسالهٔ ماست، به‌مثابهٔ نقطهٔ اتکای مقایسه‌ای عمل کرد.",
        "opt": "منطقِ «داوریِ متعادل و مشروط» در بخشِ ۵-۳ نتیجه‌گیری مستقیماً از الگوی تیلر الهام گرفت. این منبع "
               "به ما کمک کرد از قضاوتِ صفر-و-یکی پرهیز کنیم و به گزارهٔ دقیقِ «هدفِ انسجام‌بخشی درست بود، اما ابزارِ "
               "معاهدهٔ چترِ واحد در بسترِ فعلی ناکارآمد است» برسیم. هشدارِ صریحِ این منبع دربارهٔ «خطرِ تضعیفِ رژیم‌های موجود» "
               "به‌طورِ مستقیم مبنای ستونِ چهارمِ نقشهٔ راه (هم‌افزایی به‌جای افزودنِ لایهٔ هنجاریِ نو) و پاسخِ ستونِ ششم به تردیدِ ارزش‌افزوده شد.",
    },
    "DB38": {
        "points": [
            "ناکامیِ مذاکراتِ پیمان را هم‌زمان می‌توان «فرصتی ازدست‌رفته» و هم «تیرِ ازکنارگذشته» خواند.",
            "علل شکست: نبودِ ارادهٔ سیاسی، بیمِ دولت‌ها از الزام‌آورشدنِ اصول، و رقابتِ نهادی میان بازیگران.",
            "نتیجهٔ ملموسِ فرایند: قطعنامهٔ ۷۲/۲۷۷ و روندِ رسمیِ «به‌سوی یک پیمان جهانی برای محیط زیست».",
            "تحلیلِ دستِ‌اول از مواضعِ بازیگرانِ کلیدی و صف‌بندیِ ژئوپلیتیکِ مذاکرات.",
        ],
        "usage": "این یادداشتِ تحلیلی منبعِ اصلیِ تحلیلِ «علل شکست و مواضعِ ژئوپلیتیک» در فصلِ چهارم و پشتوانهٔ "
                 "بخشِ نتیجه‌گیریِ متوازن در فصلِ پنجم بود. یافته‌های آن به پرسشِ اصلی و پرسشِ فرعیِ دوم پاسخ می‌دهد.",
        "opt": "این منبع نتیجه‌گیری را از یک نگاهِ تک‌بُعدی نجات داد: به‌جای آنکه شکستِ پیمان را صرفاً «ناکامی» بدانیم، "
               "چارچوبِ دوگانهٔ «فرصتِ ازدست‌رفته / تیرِ ازکنارگذشته» را وارد کردیم و نشان دادیم فرایندِ شکست‌خورده، خود، "
               "دستاوردِ حقوقیِ ماندگار داشت. تحلیلِ سه‌گانهٔ علل شکست (حقوقی/اقتصادی/ژئوپلیتیک) و به‌ویژه «رقابتِ نهادی» "
               "مستقیماً مبنای ستونِ دومِ نقشهٔ راه (تقویتِ نهادیِ یونپ) و «بازدهِ نرمِ فرایند» مبنای ستونِ پنجم شد.",
    },
    "DB37": {
        "points": [
            "نظامِ حقوق بین‌الملل یک «نظامِ حقوقی» یکپارچه است، نه انبوهی از قواعدِ بی‌ارتباط.",
            "ابزارهای مدیریتِ تعارضِ هنجاری: قاعدهٔ خاصِ مقدم (lex specialis) و نظام‌های خودبسنده (self-contained regimes).",
            "«یکپارچه‌سازیِ نظام‌مند» از رهگذرِ مادهٔ ۳۱(۳)(ج) کنوانسیون وین، سازوکارِ کلیدیِ انسجام‌بخشی است.",
            "لنگرِ اصطلاح‌شناسیِ مستندِ رساله: همهٔ اصطلاحاتِ فنیِ تحلیلِ تعاملِ پیمان با MEAها از آن گرفته شده.",
        ],
        "usage": "نتیجه‌گیریِ رسمیِ کمیسیون حقوق بین‌الملل (۲۰۰۶، سند A/61/10 بند ۲۵۱) لنگرِ نظری و اصطلاح‌شناختیِ رساله است. "
                 "در واژه‌نامه، در فصلِ دوم و سوم (چندپارگی و ابزارهای انسجام‌بخشی) و در فصلِ چهارم (سازوکارِ حقوقیِ اثرِ پیمان بر MEAها) "
                 "به‌کار رفت و به پرسشِ اصلی و پرسشِ فرعیِ نخست پاسخ می‌دهد.",
        "opt": "این سند به نتیجه‌گیری پشتوانهٔ نظریِ رسمی بخشید. استدلالِ محوریِ رساله — اینکه اثرِ پیمان بر MEAها باید از رهگذرِ "
               "«تفسیرِ هماهنگ» و مادهٔ ۳۱(۳)(ج) وین تحلیل شود، نه صرفاً توصیف — مستقیماً از این چارچوب برخاست و به ستونِ ششمِ "
               "نقشهٔ راه (یکسان‌سازیِ تفسیرِ اصول) بدل شد. همین سند مبنای گزارهٔ کلیدیِ نتیجه‌گیری شد که چرا «معاهدهٔ چترِ واحد» "
               "با سرشتِ چندپاره و بخشیِ این نظام در تنش است.",
    },
    "DB36": {
        "points": [
            "در عصرِ انتروپوسن، معاهداتِ محیط‌زیستی به‌سببِ ساختارِ بخشی، ضعفِ اجرا و وابستگی به اجماع در «بن‌بست» گرفتار شده‌اند.",
            "حکمرانیِ سیّارهٔ زمین نیازمندِ بازاندیشیِ بنیادین در سازوکارهای الزام‌آورسازی و نظارت است.",
            "تبیینِ روشنِ چراییِ ناکارآمدیِ معاهداتِ پرشمار و پراکنده.",
            "الگوی نتیجه‌گیریِ اجرا-محور: سنتز ← دلالت‌های سیاستی ← پیشنهادِ پژوهشِ آینده.",
        ],
        "usage": "این اثر در فصلِ دوم (زمینهٔ انتروپوسن)، فصلِ چهارم و پنجم (چراییِ ناکارآمدیِ معاهدات و دلالت‌های سیاستی) به‌کار رفت "
                 "و به پرسشِ اصلی و پرسشِ فرعیِ دوم مربوط می‌شود. با توجه به ماهیتِ آن، همواره در کنارِ منابعِ داوری‌شده استناد شده است.",
        "opt": "ساختارِ «دلالت‌های سیاستی + پیشنهادِ پژوهشِ آینده» در نتیجه‌گیری از این اثر گرفته شد و تشخیصِ «ضعفِ اجرا و پایبندی» "
               "به‌مثابه ریشهٔ بن‌بست، مبنای ستونِ سومِ نقشهٔ راه (بازنگریِ ادواریِ زیست‌محیطی به الگوی UPR) شد. با این‌همه، به‌سببِ "
               "داوری‌نشده‌بودن، وزنِ استنادیِ آن تعمداً محدود و همواره در کنارِ تیلر (داوری‌شده) به‌کار رفت.",
        "caution": True,
    },
    "DB40": {
        "points": [
            "متنِ اصلیِ پیش‌نویسِ پیمان در ۲۶ ماده که اصولِ بنیادینِ حقوق بین‌الملل محیط زیست را در یک سندِ چترِ الزام‌آورِ واحد گرد می‌آورد.",
            "اصولِ کلیدی: پیشگیری، احتیاط، «آلوده‌ساز می‌پردازد»، عدالتِ بین‌نسلی، عدم‌پس‌رفت، و حقِ محیط‌زیستِ سالم.",
            "هیچ‌گاه به معاهدهٔ لازم‌الاجرا تبدیل نشد و صرفاً به قطعنامهٔ ۷۲/۲۷۷ انجامید.",
        ],
        "usage": "متنِ رسمیِ پیش‌نویس (۲۰۱۷، گروهِ حقوق‌دانان / IUCN) منبعِ دستِ‌اولِ تحلیلِ محتواییِ اصولِ پیمان در فصلِ سوم بود "
                 "و برای مقایسهٔ ماده‌به‌مادهٔ اصولِ پیمان با اصولِ مندرج در MEAها و اعلامیهٔ ریو در فصلِ چهارم به‌کار رفت. "
                 "(این سند جزوِ فایل‌های اصلیِ آپلودی نیست، اما به‌عنوانِ متنِ مبنا در تحلیل حضور دارد.)",
        "opt": "دسترسی به متنِ دقیقِ ۲۶ ماده اجازه داد نتیجه‌گیری به‌جای احکامِ کلی، بر تحلیلِ مشخصِ محتوای هر اصل استوار شود؛ "
               "و همین امکانِ سنجشِ واقع‌بینانهٔ «ارزش‌افزوده در برابرِ خطرِ تضعیفِ رژیم‌های موجود» را فراهم کرد.",
    },
}

SYNTHESIS = [
    "جمع‌بندیِ روش‌شناختی: نتیجه‌گیریِ رساله آگاهانه بر تلفیقِ سه الگوی شاخصِ هم‌موضوع بنا شده است. "
    "از تیگره (۲۰۲۲) سبکِ «استنادِ متراکمِ درون‌متنی» و ساختارِ «سنتز ← گزینه‌ها»؛ از داموآ و همکاران (۲۰۲۶، با احتیاطِ داوری‌نشده) "
    "ساختارِ «دلالت‌های سیاستی + پیشنهادِ پژوهشِ آینده»؛ و از تیلر (۲۰۱۸) منطقِ «داوریِ متعادل و مشروط». "
    "بر بسترِ سندِ پراکندگیِ کمیسیون حقوق بین‌الملل (۲۰۰۶) به‌مثابهٔ چارچوبِ نظری، این سه الگو با افزودنِ لایهٔ آینده‌پژوهی "
    "(سناریو + نقشهٔ راه) در هم آمیخته‌اند. همهٔ گزاره‌های نتیجه‌گیری یا به منبعِ معتبر یا به تحلیلِ فصولِ پیشین مستند شده‌اند.",

    "چگونه این منابع سناریوی ما را بهینه کردند؟ سناریوی اولیهٔ نگارنده پنج خلأ داشت که این منابع در رفعِ آن‌ها نقشِ مستقیم داشتند: "
    "(۱) تنشِ زمانیِ پرسشِ اصلی (فعلِ «خواهد داشت» در برابرِ واقعیتِ شکستِ ۲۰۲۲) با کمکِ چارچوبِ دوگانهٔ چاباسون و اِژه بازقاب‌بندی شد به "
    "«تأثیر بالقوه + تحلیلِ خلافِ‌واقع + نقشهٔ راه»؛ (۲) تحلیلِ تطبیقیِ شکست‌ها با اتکا به یادداشتِ IDDRI از حدِ اشاره به یک گفتارِ مستقل ارتقا یافت؛ "
    "(۳) سازوکارِ حقوقیِ اثرِ پیمان بر MEAها با تکیه بر سند کمیسیون حقوق بین‌الملل و مادهٔ ۳۱(۳)(ج) وین از توصیف به تحلیلِ هنجاری عمق گرفت؛ "
    "(۴) راهکارِ جایگزینِ سازنده (نقشهٔ راهِ تدریجیِ گذارِ اصول از حقوق نرم به سند چتر) با الهام از رویکردِ حق‌محورِ تیگره پررنگ شد؛ "
    "و (۵) پیوندِ صریح با قطعنامهٔ حق بر محیط‌زیستِ سالم (۷۶/۳۰۰) به‌مثابهٔ میراثِ زندهٔ ابتکار افزوده شد.",

    "نتیجهٔ نهایی: تلفیقِ این منابع به نتیجه‌گیریِ رساله اجازه داد از «روایتِ شکست» به «تحلیلِ دستاوردهای پنهانِ فرایند» گذر کند و "
    "به گزارهٔ متوازن و مستندِ زیر برسد: هدفِ انسجام‌بخشیِ اصول از رهگذرِ یک سندِ چتر مطلوب بود، اما ابزارِ «معاهدهٔ واحدِ اجماع‌محور» با "
    "سرشتِ واکنشی، بخشی و چندپارهٔ حقوق بین‌الملل محیط زیست ناسازگار است؛ ازاین‌رو مسیرِ واقع‌بینانه، گذارِ تدریجیِ اصول از حقوق نرم به "
    "الزامِ سخت، همراه با تقویتِ تفسیرِ هماهنگ میانِ رژیم‌ها، است.",
]

# concrete upgrades applied to the conclusion after this digest
UPGRADES_INTRO = (
    "در پیِ این هضم، به‌طورِ مشخص نتیجه‌گیریِ رساله بازبینی و تقویت شد تا هر پیشنهاد «مرجع‌محور» و همراه با "
    "«استدلالِ روشن» باشد. مهم‌ترین ارتقا، بازنویسیِ نقشهٔ راهِ شش‌ستونیِ جایگزین بود: پیش‌تر ستون‌ها فقط نام‌گذاری "
    "شده بودند، اما اکنون هر ستون به یک منبعِ معتبرِ مشخص گره خورده و توضیح می‌دهد که به کدام علتِ شکستِ پیمان پاسخ "
    "می‌دهد و چرا با سرشتِ حقوق بین‌الملل محیط زیست سازگار است:"
)
UPGRADES = [
    "ستونِ یکم (توسعهٔ رویهٔ قضایی) ← بر تزِ تیگره (۲۰۲۲) و نظرِ مشورتیِ دیوان بین‌المللی دادگستری (۲۰۲۵) استوار شد؛ پاسخ به علتِ «اجماع‌گریزی».",
    "ستونِ دوم (تقویتِ نهادیِ یونپ) ← بر تحلیلِ چاباسون و اِژه (۲۰۱۹) از «رقابتِ نهادی» و خلأِ هماهنگیِ کمیسیون حقوق بین‌الملل (۲۰۰۶).",
    "ستونِ سوم (بازنگریِ ادواریِ زیست‌محیطی به الگوی UPR) ← بر تشخیصِ «ضعفِ اجرا» نزدِ تیلر (۲۰۱۸) و داموآ و همکاران (۲۰۲۶، با احتیاط).",
    "ستونِ چهارم (خوشه‌بندی و هم‌افزایی) ← بر هشدارِ تیلر (۲۰۱۸) دربارهٔ خطرِ تضعیفِ رژیم‌های موجود و الگوی بازل-روتردام-استکهلم.",
    "ستونِ پنجم (حقوق نرم) ← بر گزارهٔ چاباسون و اِژه (۲۰۱۹) که بازدهِ واقعیِ فرایند را قطعنامه‌های ۷۲/۲۷۷ و ۷۶/۳۰۰ می‌داند و بر تزِ تیگره.",
    "ستونِ ششم (یکسان‌سازیِ تفسیر با مادهٔ ۳۱(۳)(ج) وین) ← بر سندِ کمیسیون حقوق بین‌الملل (۲۰۰۶) و در پاسخِ مستقیم به تردیدِ تیلر (۲۰۱۸).",
]
UPGRADES_TAIL = (
    "افزون بر این، بندِ پایانیِ «جمع‌بندیِ یافته‌ها» (۵-۱) با یک گزارهٔ سومِ متوازن تقویت شد — نه «ناکامیِ محض» و نه «کامیابیِ پنهان» — "
    "و بندِ «پیشنهادها و نقشهٔ راهِ عملی» (۵-۴) نیز بازنویسی شد تا هر یک از شش پیشنهاد، لنگرِ مرجعیِ خود را به‌همراه داشته باشد. "
    "این ارتقاها در نسخهٔ به‌روزِ فایلِ Word و PDF رساله و در بخشِ «نتیجه‌گیری» و «سناریو»ی سایت بازتاب یافته است. "
    "سناریوی پنج‌فصلیِ نگارنده حفظ شد و صرفاً همین ارتقاهای هدف‌مند بر آن اعمال گردید؛ نیازی به بازنویسیِ ساختارِ سناریو یا اصلاحِ بدنهٔ رساله فراتر از این نبود."
)


# ── OOXML ordering-safe paragraph decorations ──────────────────────────────
# pBdr and shd must precede bidi/spacing/jc in a valid <w:pPr>; append blindly
# and Word (though not LibreOffice) may drop them. Insert them before the first
# "later" child so both Word and LibreOffice honour the colour/border.
_LATER = {qn("w:tabs"), qn("w:bidi"), qn("w:spacing"), qn("w:ind"),
          qn("w:contextualSpacing"), qn("w:jc"), qn("w:textDirection"), qn("w:rPr")}

def _insert_ordered(p, el):
    ppr = p._p.get_or_add_pPr()
    ref = None
    for child in ppr:
        if child.tag in _LATER:
            ref = child; break
    if ref is not None:
        ref.addprevious(el)
    else:
        ppr.append(el)

def set_shading(p, fill):
    shd = OxmlElement("w:shd")
    shd.set(w("val"), "clear"); shd.set(w("color"), "auto"); shd.set(w("fill"), fill)
    _insert_ordered(p, shd)

def set_borders(p, color=NAVY, sz="6", space="4", sides=("top", "bottom", "left", "right")):
    pbdr = OxmlElement("w:pBdr")
    for s in sides:
        el = OxmlElement("w:" + s)
        el.set(w("val"), "single"); el.set(w("sz"), sz)
        el.set(w("space"), space); el.set(w("color"), color)
        pbdr.append(el)
    _insert_ordered(p, pbdr)


def hr(doc, color=GOLD, sz="6"):
    p = doc.add_paragraph()
    B.make_rtl(p, justify=False, after=2)
    set_borders(p, color=color, sz=sz, space="1", sides=("bottom",))
    return p


def heading_bar(doc, text, fill=NAVY, fg="ffffff", fa_pt=15, before=14):
    """Full-width colour bar heading — right-aligned RTL, white text."""
    p = doc.add_paragraph(); B.make_rtl(p, justify=False, before=before, after=8)
    p.paragraph_format.right_indent = Mm(1); p.paragraph_format.left_indent = Mm(1)
    set_shading(p, fill)
    set_borders(p, color=fill, sz="2", space="6")
    r = p.add_run(text); B.set_run_fonts(r, fa_pt, 12, bold=True, color=fg)
    return p


def subheading(doc, text, fa_pt=13):
    """Light band subheading with a thick gold right accent border."""
    p = doc.add_paragraph(); B.make_rtl(p, justify=False, before=10, after=6)
    p.paragraph_format.right_indent = Mm(1)
    set_shading(p, SUBBG)
    set_borders(p, color=GOLD, sz="24", space="6", sides=("right",))
    r = p.add_run(text); B.set_run_fonts(r, fa_pt, 12, bold=True, color=NAVY)
    return p


def label_line(doc, label, value, color=INK, fill=None, after=3):
    p = doc.add_paragraph(); B.make_rtl(p, justify=False, after=after)
    if fill:
        set_shading(p, fill)
        p.paragraph_format.right_indent = Mm(3); p.paragraph_format.left_indent = Mm(3)
    r = p.add_run(label + " "); B.set_run_fonts(r, 12, 11, bold=True, color=color)
    r2 = p.add_run(value); B.set_run_fonts(r2, 12, 11, bold=False)
    return p


def bullet(doc, text, mark="◄ ", mcolor=GOLD):
    p = doc.add_paragraph(); B.make_rtl(p, justify=True, after=4)
    p.paragraph_format.right_indent = Mm(6)
    r = p.add_run(mark); B.set_run_fonts(r, 13, 11, bold=True, color=mcolor)
    r2 = p.add_run(text); B.set_run_fonts(r2, 13, 11)
    return p


def banner(doc, text, fill=NAVY, fg="ffffff", fa_pt=22, before=0, after=8):
    """Centered colour banner (used on the cover)."""
    p = doc.add_paragraph(); B.make_rtl(p, justify=False, before=before, after=after)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_shading(p, fill)
    set_borders(p, color=fill, sz="2", space="10")
    r = p.add_run(text); B.set_run_fonts(r, fa_pt, 13, bold=True, color=fg)
    return p


def main():
    DB = json.load(open(os.path.join(HERE, "..", "web", "db_mining.json"), encoding="utf-8"))
    arts = DB["articles"] if isinstance(DB, dict) and "articles" in DB else DB
    by_id = {a["id"]: a for a in arts}

    doc = Document(); B.build_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Mm(297); sec.page_width = Mm(210)
    sec.top_margin = Mm(30); sec.bottom_margin = Mm(28)
    sec.right_margin = Mm(28); sec.left_margin = Mm(28)
    B.set_rtl_section(doc)
    fns = B.Footnotes(doc)

    # ── cover ──
    for _ in range(3):
        doc.add_paragraph()
    banner(doc, "گزارشِ تفصیلیِ تحلیل و هضمِ منابعِ کلیدیِ پژوهش", fill=NAVY, fa_pt=22, after=0)
    banner(doc, "شکل‌گیریِ پیمانِ جهانیِ محیط زیست و تأثیرات آن بر موافقت‌نامه‌های چندجانبهٔ محیط‌زیستی",
           fill="1c4a63", fa_pt=14, after=0)
    banner(doc, "سندِ شفافیتِ استناد: از منبع تا نتیجه‌گیری", fill=GOLD, fg="1a1205", fa_pt=13, after=6)
    for _ in range(6):
        doc.add_paragraph()
    hr(doc)
    B.center_para(doc, "تهیه‌شده توسطِ سامانهٔ هوشمندِ پژوهش (Project Brain)", fa_pt=12, after=4)
    B.center_para(doc, "HJR's Agentic Architecture @2026 — Dedicated for Dr. Masoud Ahsannejad", fa_pt=11, after=4)
    doc.add_page_break()

    # ── intro ──
    heading_bar(doc, "مقدمه: روشِ هضم و تحلیلِ منابع", fa_pt=16, before=0)
    for para in INTRO:
        B.add_para(doc, fns, para, fa_pt=14, indent=6)
    doc.add_paragraph()

    subheading(doc, "فهرستِ منابعِ بررسی‌شده")
    for i, sid in enumerate(ORDER, 1):
        a = by_id[sid]
        fa = PERSIAN[sid]
        pre = "⚠ " if NARR.get(sid, {}).get("caution") else ""
        line = f"{pre}{fa['name']} ({fa_num(a['year'])}) — {fa['title']}"
        p = doc.add_paragraph(); B.make_rtl(p, justify=False, after=4)
        p.paragraph_format.right_indent = Mm(2)
        set_borders(p, color="dfe6ee", sz="4", space="3", sides=("bottom",))
        r = p.add_run(f"{fa_num(i)}. "); B.set_run_fonts(r, 14, 12, bold=True, color=GOLD)
        r2 = p.add_run(line); B.set_run_fonts(r2, 12, 11)
    doc.add_page_break()

    # ── per-source ──
    for idx, sid in enumerate(ORDER, 1):
        a = by_id[sid]; nr = NARR[sid]; fa = PERSIAN[sid]
        heading_bar(doc, f"منبعِ {fa_num(idx)}: {fa['name']} ({fa_num(a['year'])})", fa_pt=15, before=0)

        subheading(doc, "الف) شناسنامهٔ کتاب‌شناختی")
        rows = [("عنوان (فارسی):", fa["title"]),
                ("عنوانِ اصلی (لاتین):", str(a.get("title", ""))),
                ("پدیدآور:", str(a.get("authors", ""))),
                ("سالِ انتشار:", fa_num(a.get("year", ""))),
                ("محلِ انتشار:", str(a.get("venue", ""))),
                ("پایگاه:", str(a.get("db", "")))]
        if a.get("doi"):
            rows.append(("DOI:", str(a.get("doi"))))
        rows.append(("درجهٔ ارتباط با رساله:", f"{fa_num(a.get('relevance',''))} از ۱۰۰"))
        for lbl, val in rows:
            label_line(doc, lbl, val, fill=CARDBG, after=1)
        if nr.get("caution"):
            p = doc.add_paragraph(); B.make_rtl(p, justify=True, after=4)
            set_shading(p, "fbeceb"); set_borders(p, color=RED, sz="18", space="6", sides=("right",))
            p.paragraph_format.right_indent = Mm(3); p.paragraph_format.left_indent = Mm(3)
            r = p.add_run("⚠ هشدارِ اعتبار: "); B.set_run_fonts(r, 13, 11, bold=True, color=RED)
            r2 = p.add_run("این منبع پیش‌چاپِ داوری‌نشده (non-peer-reviewed preprint) است و با احتیاط و در کنارِ منابعِ داوری‌شده استناد شده است.")
            B.set_run_fonts(r2, 13, 11)

        subheading(doc, "ب) دایجستِ کامل و نکاتِ استخراج‌شده")
        B.add_para(doc, fns, a.get("finding_fa", ""), fa_pt=14, indent=6)
        p = doc.add_paragraph(); B.make_rtl(p, justify=False, after=3)
        r = p.add_run("نکاتِ کلیدیِ استخراج‌شده:"); B.set_run_fonts(r, 13, 11, bold=True, color=NAVY)
        for pt in nr["points"]:
            bullet(doc, pt)
        if a.get("strength_fa"):
            label_line(doc, "نقطهٔ قوت:", a["strength_fa"], color=GREEN)
        if a.get("weakness_fa"):
            label_line(doc, "نقطهٔ ضعف/ملاحظه:", a["weakness_fa"], color=RED)

        subheading(doc, "ج) کجا و چگونه در پژوهش به‌کار رفت")
        label_line(doc, "فصولِ کاربرد:", "، ".join(a.get("chapters", [])) or "—", fill=CARDBG, after=1)
        label_line(doc, "پرسش‌های مرتبط:", "، ".join(a.get("questions", [])) or "—", fill=CARDBG, after=1)
        B.add_para(doc, fns, nr["usage"], fa_pt=14, indent=6)

        subheading(doc, "د) سهمِ آن در بهینه‌سازیِ نتیجه‌گیری و سناریو")
        B.add_para(doc, fns, nr["opt"], fa_pt=14, indent=6)
        if a.get("citation_note_fa"):
            label_line(doc, "یادداشتِ استناد:", a["citation_note_fa"], color=PURPLE)

        hr(doc)
        if idx < len(ORDER):
            doc.add_page_break()

    # ── synthesis ──
    doc.add_page_break()
    heading_bar(doc, "جمع‌بندی: چگونه این منابع نتیجه‌گیریِ سناریوی ما را بهینه کردند", fa_pt=16, before=0)
    for para in SYNTHESIS:
        B.add_para(doc, fns, para, fa_pt=14, indent=6)

    # ── concrete upgrades ──
    doc.add_paragraph()
    heading_bar(doc, "ارتقاهای عملیِ اعمال‌شده بر نتیجه‌گیری در پیِ این هضم", fa_pt=16)
    B.add_para(doc, fns, UPGRADES_INTRO, fa_pt=14, indent=6)
    for u in UPGRADES:
        bullet(doc, u, mark="✔ ", mcolor=GREEN)
    B.add_para(doc, fns, UPGRADES_TAIL, fa_pt=14, indent=6)

    doc.add_paragraph()
    hr(doc)
    B.center_para(doc, "HJR's Agentic Architecture @2026 — Dedicated for Dr. Masoud Ahsannejad", fa_pt=11, after=4)

    fns.finalize()
    B.enable_update_fields(doc)
    doc.save(OUT)
    print("WROTE", OUT, os.path.getsize(OUT), "bytes; footnotes:", len(fns.items))


if __name__ == "__main__":
    main()
